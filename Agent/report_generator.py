"""
报告生成模块 - 仅负责组装和保存Word报告
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import PolicySegment


class ReportGenerator:
    """报告生成器 - 仅负责组装和保存Word报告"""
    
    def __init__(self):
        self.doc = None
    
    def generate_report(self, segment: PolicySegment,
                      industry_section: str = '',
                      novelty_section: str = '',
                      meeting_section: str = '',
                      investment_section: str = '') -> Document:
        """
        生成完整报告（组装各Agent生成的Markdown部分）
        
        Args:
            segment: 政策文档
            industry_section: 行业分析部分（Markdown格式）
            novelty_section: 增量分析部分（Markdown格式）
            meeting_section: 会议分析部分（Markdown格式）
            investment_section: 投资建议部分（Markdown格式）
            
        Returns:
            Word文档对象
        """
        self.doc = Document()
        
        # 1. 标题
        self._add_title(segment)
        
        # 2. 行业分析部分（政策全文不输出到报告，仅用于输入大模型）
        if industry_section:
            try:
                self.doc.add_heading('一、行业分类分析', 1)
                self._add_markdown_text(industry_section)
            except Exception as e:
                print(f"  ⚠️ 行业分析部分生成失败: {e}")
                self.doc.add_paragraph(f"（行业分析生成失败: {e}）")
        
        # 3. 增量分析部分
        if novelty_section:
            try:
                self.doc.add_heading('二、政策增量分析', 1)
                # 处理历史政策编号替换
                novelty_section = self._replace_history_policy_numbers(novelty_section)
                self._add_markdown_text(novelty_section)
            except Exception as e:
                print(f"  ⚠️ 增量分析部分生成失败: {e}")
                self.doc.add_paragraph(f"（增量分析生成失败: {e}）")
        
        # 4. 报告系列时间对比分析部分（有同系列才显示）
        if meeting_section and meeting_section.strip() and '无同系列历史政策' not in meeting_section:
            try:
                self.doc.add_heading('三、报告系列时间对比分析', 1)
                # 处理历史政策编号替换（和增量分析一样）
                meeting_section = self._replace_history_policy_numbers(meeting_section)
                self._add_markdown_text(meeting_section)
            except Exception as e:
                print(f"  ⚠️ 报告系列分析部分生成失败: {e}")
                self.doc.add_paragraph(f"（报告系列分析生成失败: {e}）")
        
        # 5. 投资建议部分（暂时关闭）
        # if investment_section:
        #     self.doc.add_heading('四、投资建议', 1)
        #     self._add_markdown_text(investment_section)
        
        return self.doc
    
    def _add_title(self, segment: PolicySegment):
        """添加标题"""
        title = self.doc.add_heading(segment.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = self.doc.add_paragraph()
        info.add_run(f"发布时间: {segment.timestamp.strftime('%Y年%m月%d日') if segment.timestamp else 'N/A'}")
        info.add_run(" | ")
        report_series = segment.metadata.get('report_series', 'N/A')
        info.add_run(f"报告系列: {report_series}")  # ⭐ 添加报告系列
        info.add_run(" | ")
        info.add_run(f"文档ID: {segment.doc_id}")
    
    def _add_policy_content(self, content: str):
        """添加政策全文内容"""
        # 按段落分割（按换行符）
        paragraphs = content.split('\n')
        for para in paragraphs:
            para = para.strip()
            if para:
                self.doc.add_paragraph(para)
    
    def _replace_history_policy_numbers(self, text: str) -> str:
        """
        替换"历史政策1"、"历史政策2"、"同系列历史政策1"等编号为实际政策标题
        
        从文本开头的历史政策列表中提取标题，然后替换后续的编号引用
        """
        import re
        
        # 提取历史政策列表（支持两种格式）
        # 格式1：增量分析中的"#### 一、相关历史政策列表"
        # 格式2：报告系列分析中的"#### 一、同系列历史政策列表"
        history_policy_pattern = r'####\s*一、.*?历史政策列表.*?(\*\*(?:相关|同系列)?历史政策\*\*[：:].*?)(?=---|####)'
        match = re.search(history_policy_pattern, text, re.DOTALL)
        
        if not match:
            return text  # 如果没有找到历史政策列表，直接返回
        
        history_list_text = match.group(1)
        
        # 提取每个历史政策的标题（格式：1. [标题]（[时间]）或 1. [标题]（[时间]，相似度：[相似度]））
        policy_titles = {}
        # 支持两种格式：有相似度和无相似度
        policy_pattern = r'(\d+)\.\s*([^（]+)（[^）]+）'
        for match in re.finditer(policy_pattern, history_list_text):
            num = match.group(1)
            title = match.group(2).strip()
            # 支持多种编号格式
            policy_titles[f'历史政策{num}'] = title
            policy_titles[f'历史政策 {num}'] = title  # 兼容空格
            policy_titles[f'同系列历史政策{num}'] = title
            policy_titles[f'同系列历史政策 {num}'] = title  # 兼容空格
        
        # 替换文本中的"历史政策1"、"历史政策2"、"同系列历史政策1"等为实际标题
        if policy_titles:
            for num_str, title in policy_titles.items():
                # 替换为"《标题》"
                text = re.sub(
                    rf'\b{re.escape(num_str)}\b',
                    f'《{title}》',
                    text
                )
        
        return text
    
    def _add_markdown_text(self, text: str):
        """将Markdown文本添加到Word文档（清理Markdown符号，支持表格）"""
        import re
        
        # ⭐ 预处理：修复常见的LLM格式错误
        text = self._fix_llm_format_issues(text)
        
        # 清理Markdown符号
        def clean_markdown(text: str) -> str:
            # 移除加粗符号 **text** 或 *text*
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **text** -> text
            text = re.sub(r'\*([^*]+)\*', r'\1', text)  # *text* -> text
            text = re.sub(r'__([^_]+)__', r'\1', text)  # __text__ -> text
            text = re.sub(r'_([^_]+)_', r'\1', text)  # _text_ -> text
            # 移除代码块标记
            text = re.sub(r'```[\w]*\n?', '', text)  # ```code``` -> code
            text = re.sub(r'`([^`]+)`', r'\1', text)  # `code` -> code
            return text.strip()
        
        def is_table_row(line: str) -> bool:
            """判断是否为表格行"""
            return line.strip().startswith('|') and line.strip().endswith('|')
        
        def is_separator_row(line: str) -> bool:
            """判断是否为表格分隔行（如 |---|---|---|）"""
            return bool(re.match(r'^\|[\s\-:]+\|[\s\-:|]+$', line.strip()))
        
        def parse_table_row(line: str) -> list:
            """解析表格行，返回单元格列表"""
            cells = line.strip().split('|')
            # 去除首尾空元素
            cells = [c.strip() for c in cells if c.strip() or cells.index(c) not in [0, len(cells)-1]]
            return [clean_markdown(c) for c in cells]
        
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            try:
                line = lines[i].strip()
                
                if not line:
                    self.doc.add_paragraph()  # 保留空行
                    i += 1
                    continue
                
                # 检查是否为表格
                if is_table_row(line):
                    table_lines = []
                    while i < len(lines) and is_table_row(lines[i].strip()):
                        if not is_separator_row(lines[i].strip()):
                            table_lines.append(lines[i].strip())
                        i += 1
                    
                    # 解析表格并添加到Word
                    if table_lines:
                        try:
                            self._add_table(table_lines)
                        except Exception as e:
                            print(f"  ⚠️ 表格添加失败: {e}")
                            # 降级：把表格内容作为普通文本添加
                            for tl in table_lines:
                                self.doc.add_paragraph(clean_markdown(tl))
                    continue
                
                # 处理Markdown标题（按顺序检查，从多到少）
                if line.startswith('####'):
                    heading_text = line.replace('####', '').strip()
                    heading_text = clean_markdown(heading_text)
                    self.doc.add_heading(heading_text, 4)
                elif line.startswith('###'):
                    heading_text = line.replace('###', '').strip()
                    heading_text = clean_markdown(heading_text)
                    self.doc.add_heading(heading_text, 3)
                elif line.startswith('##'):
                    heading_text = line.replace('##', '').strip()
                    heading_text = clean_markdown(heading_text)
                    self.doc.add_heading(heading_text, 2)
                elif line.startswith('#'):
                    heading_text = line.replace('#', '').strip()
                    heading_text = clean_markdown(heading_text)
                    self.doc.add_heading(heading_text, 1)
                elif line.startswith('-') or (line.startswith('*') and not line.startswith('**')):
                    para_text = line[1:].strip()
                    para_text = clean_markdown(para_text)
                    self.doc.add_paragraph(para_text, style='List Bullet')
                else:
                    # 普通段落，清理Markdown符号
                    line = clean_markdown(line)
                    self.doc.add_paragraph(line)
                
                i += 1
            except Exception as e:
                print(f"  ⚠️ 处理行 {i} 失败: {e}")
                i += 1  # 跳过这一行，继续处理下一行
    
    def _add_table(self, table_lines: list):
        """将Markdown表格添加到Word文档"""
        import re
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        def clean_markdown(text: str) -> str:
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'__([^_]+)__', r'\1', text)
            text = re.sub(r'_([^_]+)_', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            return text.strip()
        
        if not table_lines:
            return
        
        # 解析所有行
        rows = []
        for line in table_lines:
            cells = line.strip().split('|')
            cells = [clean_markdown(c.strip()) for c in cells if c.strip() != '' or (cells[0] == '' and cells[-1] == '')]
            # 过滤掉因split产生的首尾空字符串
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # 确定列数（取最大列数）
        num_cols = max(len(row) for row in rows)
        num_rows = len(rows)
        
        # 创建表格
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # 填充表格内容
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    cell.text = cell_text
                    # 首行加粗（表头）
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _fix_llm_format_issues(self, text: str) -> str:
        """
        修复LLM输出的常见格式问题
        
        Args:
            text: 原始文本
            
        Returns:
            修复后的文本
        """
        import re
        
        # 1. 清理 <br> 和 <br/> 标签，替换为换行符
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        
        # 2. 清理其他HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        
        # 3. 修复用 tab 分隔的表格行（转换为 | 分隔）
        # 检测：一行中有多个 tab，且不是标准的 Markdown 表格
        lines = text.split('\n')
        fixed_lines = []
        
        for line in lines:
            stripped = line.strip()
            # 如果行中有多个 tab 但不是 | 开头，尝试转换为表格格式
            if '\t' in stripped and not stripped.startswith('|'):
                # 检查是否看起来像表格数据（有3个或以上的tab分隔内容）
                parts = stripped.split('\t')
                if len(parts) >= 3:
                    # 转换为Markdown表格格式
                    fixed_line = '| ' + ' | '.join(p.strip() for p in parts) + ' |'
                    fixed_lines.append(fixed_line)
                else:
                    fixed_lines.append(line)
            else:
                fixed_lines.append(line)
        
        text = '\n'.join(fixed_lines)
        
        # 4. 修复表格：确保表格行前后有空行，并添加分隔行
        # 检测连续的表格行并确保有正确的分隔行
        lines = text.split('\n')
        result_lines = []
        in_table = False
        table_started = False
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            is_table_row = stripped.startswith('|') and stripped.endswith('|')
            is_separator = bool(re.match(r'^\|[\s\-:]+\|', stripped))
            
            if is_table_row and not is_separator:
                if not in_table:
                    # 开始新表格
                    in_table = True
                    table_started = True
                    # 确保表格前有空行
                    if result_lines and result_lines[-1].strip():
                        result_lines.append('')
                    result_lines.append(line)
                    # 在第一行（表头）后添加分隔行
                    num_cols = stripped.count('|') - 1
                    if num_cols > 0:
                        separator = '|' + '---|' * num_cols
                        result_lines.append(separator)
                else:
                    # 继续表格
                    result_lines.append(line)
            elif is_separator and in_table:
                # 已经在表格中的分隔行，跳过（我们已经添加了）
                continue
            else:
                if in_table:
                    # 结束表格
                    in_table = False
                    table_started = False
                    # 确保表格后有空行
                    if stripped:
                        result_lines.append('')
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def save(self, filepath: str):
        """保存报告"""
        if self.doc:
            self.doc.save(filepath)
